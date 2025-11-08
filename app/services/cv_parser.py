"""
CV Parser Service
Extracts raw text from PDF, DOCX, and TXT files
"""
import PyPDF2
import pdfplumber
from docx import Document
from typing import Optional, Tuple
from pathlib import Path


class CVParser:
    """CV Parser for extracting text from various file formats"""

    @staticmethod
    def parse_pdf_pypdf2(file_path: str) -> Optional[str]:
        """
        Parse PDF using PyPDF2

        Args:
            file_path: Path to PDF file

        Returns:
            Optional[str]: Extracted text or None if failed
        """
        try:
            text_content = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text:
                        text_content.append(text)

            return '\n\n'.join(text_content) if text_content else None
        except Exception as e:
            print(f"PyPDF2 parsing failed: {str(e)}")
            return None

    @staticmethod
    def parse_pdf_pdfplumber(file_path: str) -> Optional[str]:
        """
        Parse PDF using pdfplumber (better for complex layouts)

        Args:
            file_path: Path to PDF file

        Returns:
            Optional[str]: Extracted text or None if failed
        """
        try:
            text_content = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_content.append(text)

            return '\n\n'.join(text_content) if text_content else None
        except Exception as e:
            print(f"pdfplumber parsing failed: {str(e)}")
            return None

    @staticmethod
    def parse_pdf(file_path: str) -> Tuple[Optional[str], str]:
        """
        Parse PDF using both methods and return the better result

        Args:
            file_path: Path to PDF file

        Returns:
            Tuple[Optional[str], str]: (Extracted text, method used)
        """
        # Try pdfplumber first (usually better for complex layouts)
        text_plumber = CVParser.parse_pdf_pdfplumber(file_path)

        # Try PyPDF2 as fallback
        text_pypdf2 = CVParser.parse_pdf_pypdf2(file_path)

        # Choose the result with more content
        if text_plumber and text_pypdf2:
            if len(text_plumber) >= len(text_pypdf2):
                return text_plumber, "pdfplumber"
            else:
                return text_pypdf2, "pypdf2"
        elif text_plumber:
            return text_plumber, "pdfplumber"
        elif text_pypdf2:
            return text_pypdf2, "pypdf2"
        else:
            return None, "failed"

    @staticmethod
    def parse_docx(file_path: str) -> Optional[str]:
        """
        Parse DOCX file

        Args:
            file_path: Path to DOCX file

        Returns:
            Optional[str]: Extracted text or None if failed
        """
        try:
            doc = Document(file_path)
            text_content = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_content.append(row_text)

            return '\n'.join(text_content) if text_content else None
        except Exception as e:
            print(f"DOCX parsing failed: {str(e)}")
            return None

    @staticmethod
    def parse_txt(file_path: str) -> Optional[str]:
        """
        Parse TXT file

        Args:
            file_path: Path to TXT file

        Returns:
            Optional[str]: Extracted text or None if failed
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding if UTF-8 fails
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e:
                print(f"TXT parsing failed: {str(e)}")
                return None
        except Exception as e:
            print(f"TXT parsing failed: {str(e)}")
            return None

    @staticmethod
    def parse_cv(file_path: str, file_type: str) -> Tuple[Optional[str], str]:
        """
        Parse CV file based on file type

        Args:
            file_path: Path to CV file
            file_type: Type of file (pdf, docx, txt)

        Returns:
            Tuple[Optional[str], str]: (Extracted text, method/status)
        """
        if not Path(file_path).exists():
            return None, "file_not_found"

        if file_type == "pdf":
            return CVParser.parse_pdf(file_path)
        elif file_type == "docx":
            text = CVParser.parse_docx(file_path)
            return text, "python-docx" if text else "failed"
        elif file_type == "txt":
            text = CVParser.parse_txt(file_path)
            return text, "direct_read" if text else "failed"
        else:
            return None, "unsupported_type"
